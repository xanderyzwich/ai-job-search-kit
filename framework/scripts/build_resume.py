#!/usr/bin/env python3
"""Generic resume renderer: layout only, all content from a YAML file.

Renders one docx per lane defined in the content file. Changing positioning
strategy (bullet order, lane inclusion, title/summary wording) is a content
edit, not a code change. Layout matches the recruiter-coached ATS-safe
format: Tahoma, green 12pt section headers with rule, right-tab dates,
justified 10.5pt bullets, footer for page-two continuity.

Usage:
  python3 build_resume.py                    # all lanes, content file beside script
  python3 build_resume.py ic                 # one lane
  python3 build_resume.py path/to/content.yml pc

Output files land next to the content file, named by each lane's `output`.
Every build also updates `build_record.yml` beside the content file — per
lane: build timestamp, output name, and a hash of the content file — so
"is any platform-stored copy older than the last build" is a mechanical
compare, not an act of memory.
"""
import hashlib
import sys
from datetime import datetime
from pathlib import Path

import yaml
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

GREEN = RGBColor(0x00, 0x66, 0x00)
GRAY = RGBColor(0x66, 0x66, 0x66)
FONT = "Tahoma"
RIGHT = Inches(7.0)
BODY = 10


class Builder:
    def __init__(self, content):
        self.c = content
        self.doc = Document()
        sec = self.doc.sections[0]
        sec.page_width = Inches(8.5); sec.page_height = Inches(11)
        sec.top_margin = Inches(0.6); sec.bottom_margin = Inches(0.6)
        sec.left_margin = Inches(0.75); sec.right_margin = Inches(0.75)
        normal = self.doc.styles["Normal"]
        normal.font.name = FONT; normal.font.size = Pt(BODY)
        normal.paragraph_format.space_after = Pt(3)
        normal.paragraph_format.line_spacing = 1.05
        ftr = sec.footer; ftr.is_linked_to_previous = False
        fp = ftr.paragraphs[0]; fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fr = fp.add_run(content["identity"]["footer"])
        fr.font.name = FONT; fr.font.size = Pt(8); fr.font.color.rgb = GRAY

    def runp(self, p, text, *, bold=False, size=BODY, color=None, italic=False):
        r = p.add_run(text)
        r.font.name = FONT; r.font.size = Pt(size); r.bold = bold; r.italic = italic
        if color is not None:
            r.font.color.rgb = color
        return r

    def para(self, space_after=3, space_before=0):
        p = self.doc.add_paragraph()
        p.paragraph_format.space_after = Pt(space_after)
        p.paragraph_format.space_before = Pt(space_before)
        return p

    def heading(self, text):
        p = self.para(space_before=8, space_after=3)
        self.runp(p, text, bold=True, size=12, color=GREEN)
        pPr = p._p.get_or_add_pPr()
        pbdr = OxmlElement("w:pBdr"); bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single"); bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "2"); bottom.set(qn("w:color"), "006600")
        pbdr.append(bottom); pPr.append(pbdr)

    def dateline(self, left_text, right_text):
        p = self.para(space_before=6, space_after=1)
        p.paragraph_format.tab_stops.add_tab_stop(RIGHT, WD_TAB_ALIGNMENT.RIGHT)
        self.runp(p, left_text, bold=True, size=BODY)
        self.runp(p, "\t" + right_text, bold=False, size=9)

    def bullet(self, text):
        p = self.doc.add_paragraph(style="List Bullet")
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_after = Pt(3.5)
        p.paragraph_format.line_spacing = 1.05
        run = p.add_run(text); run.font.name = FONT; run.font.size = Pt(10.5)

    def build(self, lane_key, out_dir):
        c, lane = self.c, self.c["lanes"][lane_key]
        p = self.para(space_after=1)
        self.runp(p, c["identity"]["name"], bold=True, size=19, color=GREEN)
        p = self.para(space_after=6)
        self.runp(p, c["identity"]["contact_line"], size=9.5)

        p = self.para(space_after=1)
        self.runp(p, lane["title_line"], bold=True, size=12, color=GREEN)
        p = self.para(space_after=5)
        self.runp(p, lane["tagline"], bold=True, size=10)
        p = self.para(space_after=2)
        self.runp(p, lane["summary"], size=BODY)

        self.heading("Core Competencies")
        p = self.para(space_after=3)
        self.runp(p, lane["competencies"], bold=True, size=10)

        self.heading("Technical Skills")
        for label, value in c["skills"]:
            p = self.para(space_after=2)
            p.paragraph_format.line_spacing = 1.05
            self.runp(p, label + ": ", bold=True, size=BODY)
            self.runp(p, value, size=BODY)

        self.heading("Professional Experience")
        for emp in c["employers"]:
            self.dateline(emp["heading"], emp["dates"])
            p = self.para(space_after=2)
            scope = (emp.get("scope") or {}).get(lane_key)
            if scope:
                self.runp(p, emp["title"] + " ", bold=True, size=BODY)
                self.runp(p, scope, size=BODY, italic=True)
            else:
                self.runp(p, emp["title"], bold=True, size=BODY)
            context = (emp.get("context") or {}).get(lane_key)
            if context:
                p = self.para(space_after=2)
                self.runp(p, context, size=BODY, italic=True)
            order = lane.get("orders", {}).get(emp["key"]) or emp["order"]
            for key in order:
                self.bullet(emp["bullets"][key])

        self.heading("Community & Talks")
        for item in c["community"]:
            self.bullet(item)

        self.heading("Education")
        p = self.para(space_after=1)
        self.runp(p, c["education"], size=BODY)

        out = out_dir / lane["output"]
        self.doc.save(out)
        return out


def write_build_record(content_path, built):
    """Update build_record.yml beside the content file. `built` maps lane
    key -> output Path for the lanes rendered in this run; lanes not
    rebuilt keep their previous record entry."""
    record_path = content_path.parent / "build_record.yml"
    record = {}
    if record_path.exists():
        record = yaml.safe_load(record_path.read_text(encoding="utf-8")) or {}
    content_hash = hashlib.sha256(content_path.read_bytes()).hexdigest()[:16]
    for lane_key, out in built.items():
        record[lane_key] = {
            "built": datetime.now().isoformat(timespec="seconds"),
            "output": out.name,
            "content_sha256_16": content_hash,
        }
    record_path.write_text(yaml.safe_dump(record, sort_keys=False),
                           encoding="utf-8")
    return record_path


def main():
    args = sys.argv[1:]
    content_path = Path(__file__).resolve().parent / "resume_content.yml"
    lanes = None
    for a in args:
        if a.endswith((".yml", ".yaml")):
            content_path = Path(a).resolve()
        else:
            lanes = [a]
    with open(content_path, encoding="utf-8") as f:
        content = yaml.safe_load(f)
    built = {}
    for lane_key in lanes or list(content["lanes"]):
        if lane_key not in content["lanes"]:
            sys.exit(f"Unknown lane '{lane_key}'. "
                     f"Available: {', '.join(content['lanes'])}")
        out = Builder(content).build(lane_key, content_path.parent)
        built[lane_key] = out
        print("SAVED", out)
    record = write_build_record(content_path, built)
    print("RECORDED", record)


if __name__ == "__main__":
    main()
